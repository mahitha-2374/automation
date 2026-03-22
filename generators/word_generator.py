"""
Word OLA Generator
Fills OLA Word template with processed data
"""

from docx import Document
import os


class WordGenerator:
    """Generate OLA Word documents from template"""

    def __init__(self):
        """Initialize Word generator"""
        pass

    def generate(self, template_path, output_path, data):
        """
        Fill OLA template with processed data
        
        Args:
            template_path: Path to OLA_template.docx
            output_path: Where to save output
            data: Processed IAM data
        """
        if not os.path.exists(template_path):
            # Create basic template if it doesn't exist
            self._create_basic_template(template_path)

        doc = Document(template_path)

        # Prepare replacement data
        replacements = self._prepare_replacements(data)

        # Replace text in paragraphs
        self._replace_in_paragraphs(doc, replacements)

        # Fill tables if they exist
        self._fill_tables(doc, data)

        doc.save(output_path)

    def _prepare_replacements(self, data):
        """Prepare key-value replacements"""
        gsi_id = data.get("gsi_id", "NOT_PROVIDED")
        total_users = len(data.get("users", []))
        total_roles = len(data.get("roles", []))
        total_entitlements = len(data.get("entitlements", []))
        total_user_role_mappings = sum(
            len(u["roles"]) for u in data.get("users", [])
        )

        return {
            "<<GSI_ID>>": gsi_id,
            "{{GSI_ID}}": gsi_id,
            "<<SYSTEM_ID>>": gsi_id,
            "<<SYSTEM_NAME>>": f"IAM System ({gsi_id})",
            "{{SYSTEM_ID}}": gsi_id,
            "{{SYSTEM_NAME}}": f"IAM System ({gsi_id})",
            "<<USER_COUNT>>": str(total_users),
            "<<ROLE_COUNT>>": str(total_roles),
            "<<ENT_COUNT>>": str(total_entitlements),
            "<<MAPPING_COUNT>>": str(total_user_role_mappings),
            "{{USER_COUNT}}": str(total_users),
            "{{ROLE_COUNT}}": str(total_roles),
            "{{ENT_COUNT}}": str(total_entitlements),
            "{{MAPPING_COUNT}}": str(total_user_role_mappings),
        }

    def _replace_in_paragraphs(self, doc, replacements):
        """Replace text in all paragraphs"""
        for p in doc.paragraphs:
            for key, value in replacements.items():
                if key in p.text:
                    p.text = p.text.replace(key, str(value))

    def _fill_tables(self, doc, data):
        """Fill tables with data if they exist"""
        if len(doc.tables) > 0:
            # Fill summary table
            if len(doc.tables) > 0:
                table = doc.tables[0]
                if len(table.rows) > 1:
                    if len(table.columns) > 1:
                        table.cell(1, 0).text = str(len(data.get("users", [])))
                        if len(table.rows) > 2 and len(table.columns) > 1:
                            table.cell(2, 0).text = str(len(data.get("roles", [])))

    def _create_basic_template(self, path):
        """Create a basic OLA template"""
        doc = Document()

        doc.add_heading("IAM System OLA", 0)
        doc.add_paragraph("System ID: <<SYSTEM_ID>>")
        doc.add_paragraph("System Name: <<SYSTEM_NAME>>")

        doc.add_heading("Summary", 1)
        doc.add_paragraph(f"Total Users: <<USER_COUNT>>")
        doc.add_paragraph(f"Total Roles: <<ROLE_COUNT>>")
        doc.add_paragraph(f"Total Entitlements: <<ENT_COUNT>>")

        doc.save(path)
