"""
Universal Template Filler Engine
Fills ANY template (Excel, Word, PDF) with data from CSV
Supports automatic placeholder detection and intelligent data mapping
"""

import os
import re
import pandas as pd
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Any, Tuple, Optional
from copy import copy

# Import template handlers
try:
    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter
except ImportError:
    load_workbook = None

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
except ImportError:
    SimpleDocTemplate = None


class PlaceholderDetector:
    """Detects and extracts placeholders from templates"""
    
    PLACEHOLDER_PATTERNS = [
        r'\{\{(\w+)\}\}',      # {{field}}
        r'\[(\w+)\]',          # [FIELD]
        r'\{(\w+)\}',          # {field} - generic
        r'\$(\w+)\$',          # $field$ - alternate
    ]
    
    @classmethod
    def detect_placeholders(cls, text: str) -> List[str]:
        """Extract all placeholders from text"""
        placeholders = set()
        for pattern in cls.PLACEHOLDER_PATTERNS:
            matches = re.findall(pattern, text, re.IGNORECASE)
            placeholders.update(matches)
        return list(placeholders)
    
    @classmethod
    def replace_placeholder(cls, text: str, field_name: str, value: Any) -> str:
        """Replace all variations of a placeholder with value"""
        value_str = str(value) if value is not None else ""
        
        # Replace all placeholder patterns
        text = re.sub(r'\{\{' + field_name + r'\}\}', value_str, text, flags=re.IGNORECASE)
        text = re.sub(r'\[' + field_name + r'\]', value_str, text, flags=re.IGNORECASE)
        text = re.sub(r'\{' + field_name + r'\}', value_str, text, flags=re.IGNORECASE)
        text = re.sub(r'\$' + field_name + r'\$', value_str, text, flags=re.IGNORECASE)
        
        return text


class DataMapper:
    """Intelligently maps CSV data to template fields"""
    
    # Field name variations that map to the same concept
    FIELD_ALIASES = {
        'user': ['user', 'userid', 'user_id', 'username', 'user_name', 'login'],
        'role': ['role', 'roleid', 'role_id', 'role_name'],
        'gsi': ['gsi', 'gsi_id', 'gsid', 'system', 'app', 'application'],
        'status': ['status', 'account_status', 'accountstatus', 'active'],
        'manager': ['manager', 'manager_id', 'managerid', 'supervisor'],
        'description': ['description', 'desc', 'des'],
        'resource': ['resource', 'resname', 'res_name', 'resource_name'],
        'module': ['module', 'module_name', 'modulename'],
    }
    
    @classmethod
    def normalize_field_name(cls, field: str) -> str:
        """Normalize field name to standard form"""
        field_lower = field.lower()
        for standard, aliases in cls.FIELD_ALIASES.items():
            if field_lower in aliases:
                return standard
        return field_lower
    
    @classmethod
    def find_matching_column(cls, placeholder: str, df: pd.DataFrame) -> Optional[str]:
        """Find actual column in dataframe that matches placeholder"""
        normalized_placeholder = cls.normalize_field_name(placeholder)
        
        for col in df.columns:
            if cls.normalize_field_name(col) == normalized_placeholder:
                return col
        
        return None
    
    @classmethod
    def get_summary_stats(cls, users_df: pd.DataFrame, roles_df: pd.DataFrame, 
                         entitlements: List[Dict]) -> Dict[str, Any]:
        """Generate summary statistics from data"""
        return {
            'total_users': len(users_df),
            'total_roles': len(roles_df),
            'total_entitlements': len(entitlements),
            'unique_users': users_df['UserID'].nunique() if 'UserID' in users_df.columns else 0,
            'unique_roles': roles_df['Role_Name'].nunique() if 'Role_Name' in roles_df.columns else 0,
            'export_date': datetime.now().strftime('%Y-%m-%d'),
            'export_time': datetime.now().strftime('%H:%M:%S'),
        }


class ExcelTemplateFiller:
    """Fills Excel templates with data"""
    
    @staticmethod
    def fill_template(template_path: str, output_path: str, data: Dict[str, Any]) -> bool:
        """Fill Excel template with data"""
        if not load_workbook:
            print("⚠️ openpyxl not available")
            return False
        
        try:
            wb = load_workbook(template_path)
            
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                
                # Fill cell values
                for row in ws.iter_rows():
                    for cell in row:
                        if cell.value and isinstance(cell.value, str):
                            # Detect and replace placeholders
                            placeholders = PlaceholderDetector.detect_placeholders(cell.value)
                            for ph in placeholders:
                                col = DataMapper.find_matching_column(ph, pd.DataFrame([data]))
                                if col and col in data:
                                    cell.value = PlaceholderDetector.replace_placeholder(
                                        cell.value, ph, data.get(ph, data.get(col, ""))
                                    )
                
                # Handle range-based data (for multiple rows)
                # Look for markers like {{users_start}} and {{users_end}}
                ExcelTemplateFiller._fill_data_ranges(ws, data)
            
            wb.save(output_path)
            return True
        except Exception as e:
            print(f"❌ Error filling Excel template: {str(e)}")
            return False
    
    @staticmethod
    def _fill_data_ranges(ws, data: Dict[str, Any]):
        """Fill data ranges (e.g., user lists, role lists)"""
        try:
            # Find markers for data ranges
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        if '{{users_start}}' in cell.value or '[USERS_START]' in cell.value:
                            # Mark this row for user data
                            if 'users' in data and isinstance(data['users'], list):
                                ExcelTemplateFiller._fill_rows(ws, row[0].row, data['users'])
                        
                        if '{{roles_start}}' in cell.value or '[ROLES_START]' in cell.value:
                            # Mark this row for role data
                            if 'roles' in data and isinstance(data['roles'], list):
                                ExcelTemplateFiller._fill_rows(ws, row[0].row, data['roles'])
        except Exception as e:
            pass  # Range filling is optional
    
    @staticmethod
    def _fill_rows(ws, start_row: int, data_list: List[Dict]):
        """Fill multiple rows with data"""
        try:
            for idx, data_row in enumerate(data_list):
                row_num = start_row + idx
                for col_idx, (key, value) in enumerate(data_row.items(), 1):
                    cell = ws.cell(row=row_num, column=col_idx)
                    cell.value = value
        except Exception as e:
            pass


class WordTemplateFiller:
    """Fills Word templates with data"""
    
    @staticmethod
    def fill_template(template_path: str, output_path: str, data: Dict[str, Any]) -> bool:
        """Fill Word template with data"""
        if not Document:
            print("⚠️ python-docx not available")
            return False
        
        try:
            doc = Document(template_path)
            
            # Fill paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    placeholders = PlaceholderDetector.detect_placeholders(paragraph.text)
                    for ph in placeholders:
                        # Try to find matching data
                        col = DataMapper.find_matching_column(ph, pd.DataFrame([data]))
                        value = data.get(ph, data.get(col, ""))
                        
                        # Replace in paragraph
                        for run in paragraph.runs:
                            if ph_text := PlaceholderDetector.detect_placeholders(run.text):
                                run.text = PlaceholderDetector.replace_placeholder(run.text, ph, value)
                        
                        # Also try full paragraph replacement
                        paragraph.text = PlaceholderDetector.replace_placeholder(paragraph.text, ph, value)
            
            # Fill tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                placeholders = PlaceholderDetector.detect_placeholders(paragraph.text)
                                for ph in placeholders:
                                    col = DataMapper.find_matching_column(ph, pd.DataFrame([data]))
                                    value = data.get(ph, data.get(col, ""))
                                    paragraph.text = PlaceholderDetector.replace_placeholder(
                                        paragraph.text, ph, value
                                    )
            
            doc.save(output_path)
            return True
        except Exception as e:
            print(f"❌ Error filling Word template: {str(e)}")
            return False


class UniversalTemplateFiller:
    """Main template filler - handles all formats"""
    
    SUPPORTED_FORMATS = {
        '.xlsx': 'Excel',
        '.xlsm': 'Excel',
        '.docx': 'Word',
        '.pdf': 'PDF'
    }
    
    def __init__(self, template_dir: str = "templates"):
        self.template_dir = template_dir
        self.stats = {}
    
    def fill_template(self, template_name: str, output_path: str, 
                     users_df: pd.DataFrame, roles_df: pd.DataFrame,
                     gsi_id: Optional[str] = None) -> Tuple[bool, str]:
        """
        Fill any template with data
        
        Returns: (success, message)
        """
        # Find template file
        template_files = self._find_template(template_name)
        if not template_files:
            return False, f"Template '{template_name}' not found"
        
        template_path = template_files[0]
        ext = Path(template_path).suffix.lower()
        
        # Prepare data
        data = self._prepare_data(users_df, roles_df, gsi_id)
        
        # Fill based on format
        success = False
        if ext in ['.xlsx', '.xlsm']:
            success = ExcelTemplateFiller.fill_template(template_path, output_path, data)
        elif ext == '.docx':
            success = WordTemplateFiller.fill_template(template_path, output_path, data)
        else:
            return False, f"Unsupported format: {ext}"
        
        if success:
            return True, f"✓ {template_name} generated: {Path(output_path).name}"
        else:
            return False, f"✗ Failed to fill {template_name}"
    
    def _find_template(self, template_name: str) -> List[str]:
        """Find template file by name"""
        if not os.path.exists(self.template_dir):
            return []
        
        matches = []
        for file in os.listdir(self.template_dir):
            if template_name.lower() in file.lower():
                matches.append(os.path.join(self.template_dir, file))
        
        return matches
    
    def _prepare_data(self, users_df: pd.DataFrame, roles_df: pd.DataFrame,
                     gsi_id: Optional[str] = None) -> Dict[str, Any]:
        """Prepare data dictionary from dataframes"""
        # Filter by GSI if provided
        if gsi_id:
            # Check for gsi column
            gsi_col = self._find_gsi_column(users_df)
            if gsi_col:
                users_df = users_df[users_df[gsi_col] == gsi_id]
                # Also try to filter roles
                if self._find_gsi_column(roles_df):
                    roles_df = roles_df[roles_df[self._find_gsi_column(roles_df)] == gsi_id]
        
        # Create data dictionary
        data = {
            'gsi_id': gsi_id or 'UNKNOWN',
            'total_users': len(users_df),
            'total_roles': len(roles_df),
            'export_date': datetime.now().strftime('%Y-%m-%d'),
            'export_time': datetime.now().strftime('%H:%M:%S'),
            'users': users_df.to_dict('records'),
            'roles': roles_df.to_dict('records'),
        }
        
        # Add summary stats
        data.update(DataMapper.get_summary_stats(users_df, roles_df, []))
        
        return data
    
    def _find_gsi_column(self, df: pd.DataFrame) -> Optional[str]:
        """Find GSI column in dataframe"""
        for col in df.columns:
            if col.lower() in ['gsi', 'gsi_id', 'system', 'app', 'application']:
                return col
        return None
    
    def get_available_templates(self) -> Dict[str, Dict]:
        """List all available templates"""
        templates = {}
        
        if not os.path.exists(self.template_dir):
            return templates
        
        for file in os.listdir(self.template_dir):
            ext = Path(file).suffix.lower()
            if ext in self.SUPPORTED_FORMATS:
                template_name = Path(file).stem
                templates[template_name] = {
                    'file': file,
                    'format': ext,
                    'type': self.SUPPORTED_FORMATS.get(ext, 'Unknown'),
                    'path': os.path.join(self.template_dir, file)
                }
        
        return templates
